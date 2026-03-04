from datetime import datetime

# import comtypes.client
from flask import Flask, render_template, request, redirect, url_for, send_from_directory
import os
from utils.converter import (
    convert_pdf_to_docx,
    convert_docx_to_pdf,
    convert_docx_to_html,
    convert_word_to_pdf,
    convert_ppt_to_pdf,
    convert_excel_to_pdf
)
from dotenv import load_dotenv
load_dotenv()
from database import init_db, get_connection
PDF_STORAGE = {}

app = Flask(__name__)
init_db()

UPLOAD_FOLDER = "static/uploads"
CONVERTED_FOLDER = "static/converted"

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["CONVERTED_FOLDER"] = CONVERTED_FOLDER

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CONVERTED_FOLDER, exist_ok=True)

@app.route("/")
def home():
    return render_template("index.html")


@app.route("/upload", methods=["POST"])
def upload_file():

    if "pdf_files" not in request.files:
        return "No files uploaded"

    files = request.files.getlist("pdf_files")

    if not files or files[0].filename == "":
        return "No selected files"

    for file in files:
        if file and file.filename.endswith(".pdf"):

            upload_path = os.path.join(
                app.config["UPLOAD_FOLDER"],
                file.filename
            )
            file.save(upload_path)

            # 🔥 Convert PDF → DOCX
            docx_filename = file.filename.replace(".pdf", ".docx")
            docx_path = os.path.join(
                app.config["CONVERTED_FOLDER"],
                docx_filename
            )

            convert_pdf_to_docx(upload_path, docx_path)

            # 🔥 Redirect to Word-style editor
            return redirect(url_for(
                "editor",
                filename=docx_filename
            ))

    return "Invalid file type"


@app.route("/editor/<filename>")
def editor(filename):
    docx_path = os.path.join(app.config["CONVERTED_FOLDER"], filename)
    
    html_content = convert_docx_to_html(docx_path)

    return render_template("editor.html", 
                           filename=filename, 
                           content=html_content)

@app.route("/documents")
def documents():
    files = []

    for filename in os.listdir(app.config["CONVERTED_FOLDER"]):
        if filename.endswith(".docx"):
            file_path = os.path.join(app.config["CONVERTED_FOLDER"], filename)
            modified_time = os.path.getmtime(file_path)

            files.append({
                "name": filename,
                "modified": modified_time
            })

    # Sort newest first
    files.sort(key=lambda x: x["modified"], reverse=True)

    return render_template("documents.html", files=files)

@app.route("/delete/<filename>")
def delete_file(filename):
    file_path = os.path.join(app.config["CONVERTED_FOLDER"], filename)

    if os.path.exists(file_path):
        os.remove(file_path)

    return redirect(url_for("documents"))

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn



def html_to_docx(html_content, output_path, header_text=None, footer_text=None, watermark_enabled=False, watermark_text=None):
    soup = BeautifulSoup(html_content, "html.parser")
    doc = Document()
    # =====================
    # WATERMARK
    # =====================
    if watermark_enabled and watermark_text:

        watermark_xml = f'''
        <w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                xmlns:v="urn:schemas-microsoft-com:vml">
        <v:shape id="PowerPlusWaterMarkObject"
            style="position:absolute;
                    width:468pt;
                    height:117pt;
                    rotation:315;
                    z-index:-251654144"
            fillcolor="silver"
            stroked="f">
            <v:textpath style="font-family:'Calibri';font-size:48pt"
                string="{watermark_text}"/>
        </v:shape>
        </w:pict>
        '''

        watermark = parse_xml(watermark_xml)
        doc.sections[0].header._element.append(watermark)
    

    section = doc.sections[0]

    if header_text:
        header = section.header
        header.paragraphs[0].text = header_text

    if footer_text:
        footer = section.footer
        footer.paragraphs[0].text = footer_text
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    container = soup.body if soup.body else soup

    from bs4 import NavigableString

    for element in container.contents:

        # Skip empty text nodes
        if isinstance(element, NavigableString):
            if str(element).strip():
                doc.add_paragraph(str(element).strip())
            continue

        # =====================
        # HEADINGS
        # =====================
        if element.name in ["h1", "h2", "h3"]:
            level = int(element.name[1])
            heading = doc.add_heading(element.get_text(), level=level)
            heading.paragraph_format.space_after = Pt(12)
            continue

        # =====================
        # PARAGRAPH
        # =====================
        if element.name == "p":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.5

            style = element.get("style", "")
            if "text-align: center" in style:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif "text-align: right" in style:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif "text-align: justify" in style:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            for child in element.children:

                if isinstance(child, NavigableString):
                    paragraph.add_run(str(child))
                    continue

                run = paragraph.add_run(child.get_text())

                if child.name in ["strong", "b"]:
                    run.bold = True

                if child.name in ["em", "i"]:
                    run.italic = True

                if child.name == "u":
                    run.underline = True

                if child.attrs.get("style"):
                    style = child.attrs["style"]

                    if "font-size" in style:
                        size = style.split("font-size:")[1].split("px")[0].strip()
                        run.font.size = Pt(float(size) * 0.75)

                    if "color" in style:
                        color = style.split("color:")[1].split(";")[0].strip()
                        if color.startswith("#"):
                            run.font.color.rgb = RGBColor.from_string(color[1:])

                    if "font-family" in style:
                        font_name = style.split("font-family:")[1].split(";")[0].strip()
                        run.font.name = font_name        
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

            continue

        # =====================
        # BULLET LIST
        # =====================
        if element.name == "ul":
            for li in element.find_all("li"):
                p = doc.add_paragraph(li.get_text(), style="List Bullet")
                p.paragraph_format.space_after = Pt(4)
            continue

        # =====================
        # NUMBERED LIST
        # =====================
        if element.name == "ol":
            for li in element.find_all("li"):
                p = doc.add_paragraph(li.get_text(), style="List Number")
                p.paragraph_format.space_after = Pt(4)
            continue

        # =====================
        # TABLE
        # =====================
        if element.name == "table":
            rows = element.find_all("tr")
            if not rows:
                continue

            cols = rows[0].find_all(["td", "th"])
            table = doc.add_table(rows=len(rows), cols=len(cols))
            table.style = "Table Grid"

            for i, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                for j, cell in enumerate(cells):
                    table.rows[i].cells[j].text = cell.get_text()

            continue

        # =====================
        # LINK
        # =====================
        if element.name == "a":
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(element.get_text())
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.underline = True
            continue

        # =========================
        # GLOBAL DOCUMENT STYLING
        # =========================

        style = doc.styles['Normal']
        style.font.name = "Calibri"
        style.font.size = Pt(12)

        section = doc.sections[0]
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    doc.save(output_path)

@app.route("/save/<filename>", methods=["POST"])
def save_file(filename):

    html_content = request.form.get("content")
    header_text = request.form.get("header_text")
    footer_text = request.form.get("footer_text")

    # 🔥 NEW
    watermark_enabled = request.form.get("watermark_enabled")
    watermark_text = request.form.get("watermark_text")

    output_path = os.path.join(app.config["CONVERTED_FOLDER"], filename)

    html_to_docx(
        html_content,
        output_path,
        header_text,
        footer_text,
        watermark_enabled,
        watermark_text
    )

    return redirect(url_for("editor", filename=filename))

@app.route("/autosave/<filename>", methods=["POST"])
def autosave_file(filename):
    html_content = request.form.get("content")

    if not html_content:
        return {"status": "error"}

    output_path = os.path.join(app.config["CONVERTED_FOLDER"], filename)

    html_to_docx(html_content, output_path)

    return {"status": "saved"}


@app.route("/download_pdf/<filename>")
def download_pdf(filename):
    docx_path = os.path.join(app.config["CONVERTED_FOLDER"], filename)
    
    convert_docx_to_pdf(docx_path, app.config["CONVERTED_FOLDER"])

    pdf_filename = filename.replace(".docx", ".pdf") 
    
    return send_from_directory(app.config["CONVERTED_FOLDER"], pdf_filename, as_attachment=True)

@app.route("/download/<filename>")
def download_file(filename):
    return send_from_directory(app.config["CONVERTED_FOLDER"], filename, as_attachment=True)

@app.route("/upload_page")
def upload_page():
    return render_template("upload.html")

@app.route("/dashboard")
def dashboard():
    converted_files = os.listdir(app.config["CONVERTED_FOLDER"])
    upload_files = os.listdir(app.config["UPLOAD_FOLDER"])

    total_docs = len([f for f in converted_files if f.endswith(".docx")])
    total_uploads = len([f for f in upload_files if f.endswith(".pdf")])

    total_storage = 0
    for folder in [app.config["UPLOAD_FOLDER"], app.config["CONVERTED_FOLDER"]]:
        for file in os.listdir(folder):
            file_path = os.path.join(folder, file)
            total_storage += os.path.getsize(file_path)

    # Convert to MB
    total_storage_mb = round(total_storage / (1024 * 1024), 2)

    return render_template(
        "dashboard.html",
        total_docs=total_docs,
        total_uploads=total_uploads,
        total_storage=total_storage_mb
    )

from PyPDF2 import PdfMerger

@app.route("/merge", methods=["GET", "POST"])
def merge_pdf():
    if request.method == "POST":

        files = request.files.getlist("pdf_files")
        order = request.form.get("file_order")

        merger = PdfMerger()
        output_path = os.path.join(app.config["CONVERTED_FOLDER"], "merged.pdf")

        # Convert order string to list of integers
        if order:
            order_indexes = list(map(int, order.split(",")))
        else:
            order_indexes = list(range(len(files)))

        # Merge in reordered sequence
        for index in order_indexes:
            merger.append(files[index])

        merger.write(output_path)
        merger.close()

        return send_from_directory(app.config["CONVERTED_FOLDER"],
                                   "merged.pdf",
                                   as_attachment=True)

    return render_template("merge.html")

from PyPDF2 import PdfReader, PdfWriter
import zipfile

@app.route("/split", methods=["GET", "POST"])
def split_pdf():
    if request.method == "POST":

        if "pdf_file" not in request.files:
            return "No file uploaded"

        file = request.files["pdf_file"]
        selected_pages = request.form.get("selected_pages")

        if not selected_pages:
            return "No pages selected"

        reader = PdfReader(file)
        writer = PdfWriter()

        pages = list(map(int, selected_pages.split(",")))

        for page_num in pages:
            if 0 <= page_num < len(reader.pages):
                writer.add_page(reader.pages[page_num])

        output_path = os.path.join(app.config["CONVERTED_FOLDER"], "extracted.pdf")

        with open(output_path, "wb") as f:
            writer.write(f)

        return send_from_directory(app.config["CONVERTED_FOLDER"],
                                   "extracted.pdf",
                                   as_attachment=True)

    return render_template("split.html")

import pikepdf

@app.route("/compress", methods=["GET", "POST"])
def compress_pdf():
    if request.method == "POST":
        file = request.files["pdf_file"]
        level = request.form.get("compression_level")

        input_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
        output_filename = "compressed.pdf"
        output_path = os.path.join(app.config["CONVERTED_FOLDER"], output_filename)

        file.save(input_path)

        original_size = os.path.getsize(input_path)

        with pikepdf.open(input_path) as pdf:

            if level == "low":
                pdf.save(output_path)

            elif level == "medium":
                pdf.save(output_path, compress_streams=True)

            elif level == "high":
                pdf.save(output_path,
                         compress_streams=True,
                         object_stream_mode=pikepdf.ObjectStreamMode.generate)

        compressed_size = os.path.getsize(output_path)

        original_mb = round(original_size / (1024 * 1024), 2)
        compressed_mb = round(compressed_size / (1024 * 1024), 2)

        reduction = round(((original_size - compressed_size) / original_size) * 100, 2)

        return render_template(
            "compress.html",
            original_size=original_mb,
            compressed_size=compressed_mb,
            reduction=reduction,
            download_file=output_filename
        )

    return render_template("compress.html")

# Additional routes for other features can be added here
@app.route("/word-to-pdf", methods=["GET", "POST"])
def word_to_pdf():
    if request.method == "POST":

        file = request.files.get("file")

        if not file:
            return "No file uploaded"

        input_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
        file.save(input_path)

        # Convert
        output_path = os.path.join(
            app.config["CONVERTED_FOLDER"],
            file.filename.replace(".docx", ".pdf")
        )

        convert_word_to_pdf(input_path, output_path)

        output_filename = file.filename.replace(".docx", ".pdf")

        return send_from_directory(app.config["CONVERTED_FOLDER"],
                                   output_filename,
                                   as_attachment=True)

    return render_template("word_to_pdf.html")

@app.route("/pdf-to-word", methods=["GET", "POST"])
def pdf_to_word():
    if request.method == "POST":

        file = request.files.get("file")

        if not file:
            return "No file uploaded"

        if not file.filename.lower().endswith(".pdf"):
            return "Only PDF files allowed"

        # Save uploaded file
        input_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
        file.save(input_path)

        # Convert to DOCX
        output_filename = file.filename.replace(".pdf", ".docx")
        output_path = os.path.join(app.config["CONVERTED_FOLDER"], output_filename)

        convert_pdf_to_docx(input_path, output_path)

        return send_from_directory(
            app.config["CONVERTED_FOLDER"],
            output_filename,
            as_attachment=True
        )

    return render_template("pdf_to_word.html")

from pdf2image import convert_from_path
from PyPDF2 import PdfReader
from PIL import Image, ImageDraw, ImageFont
import zipfile

@app.route("/pdf-to-jpg", methods=["GET", "POST"])
def pdf_to_jpg():
    if request.method == "POST":

        file = request.files.get("file")
        dpi = int(request.form.get("dpi", 150))
        quality = int(request.form.get("quality", 85))
        grayscale = request.form.get("grayscale")
        watermark = request.form.get("watermark")
        selected_pages = request.form.get("selected_pages")

        if not file:
            return "No file uploaded"

        input_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
        file.save(input_path)

        reader = PdfReader(input_path)
        total_pages = len(reader.pages)

        # Determine selected pages
        if selected_pages:
            page_numbers = list(map(int, selected_pages.split(",")))
        else:
            page_numbers = list(range(total_pages))

        images = convert_from_path(
            input_path,
            dpi=dpi,
            poppler_path=r"C:\poppler\Library\bin"
        )

        output_files = []

        for i in page_numbers:
            image = images[i]

            # Apply grayscale
            if grayscale:
                image = image.convert("L").convert("RGB")

            # Apply watermark
            if watermark:
                draw = ImageDraw.Draw(image)
                draw.text((50, 50), watermark, fill=(255, 0, 0))

            filename = f"page_{i+1}.jpg"
            output_path = os.path.join(app.config["CONVERTED_FOLDER"], filename)

            image.save(output_path, "JPEG", quality=quality)
            output_files.append(filename)

        return render_template(
            "pdf_to_jpg_result.html",
            files=output_files
        )

    return render_template("pdf_to_jpg.html")

@app.route("/download-all", methods=["POST"])
def download_all():
    files = request.form.get("files").split(",")

    zip_path = os.path.join(app.config["CONVERTED_FOLDER"], "all_images.zip")

    with zipfile.ZipFile(zip_path, "w") as zipf:
        for file in files:
            file_path = os.path.join(app.config["CONVERTED_FOLDER"], file)
            zipf.write(file_path, file)

    return send_from_directory(app.config["CONVERTED_FOLDER"],
                               "all_images.zip",
                               as_attachment=True)

from PIL import Image
from PyPDF2 import PdfMerger

@app.route("/jpg-to-pdf", methods=["GET", "POST"])
def jpg_to_pdf():
    if request.method == "POST":

        files = request.files.getlist("files")

        if not files:
            return "No images uploaded"

        image_list = []

        for file in files:
            image = Image.open(file).convert("RGB")
            image_list.append(image)

        output_path = os.path.join(app.config["CONVERTED_FOLDER"], "converted.pdf")

        image_list[0].save(
            output_path,
            save_all=True,
            append_images=image_list[1:]
        )

        return send_from_directory(
            app.config["CONVERTED_FOLDER"],
            "converted.pdf",
            as_attachment=True
        )

    return render_template("jpg_to_pdf.html")

@app.route("/excel-to-pdf", methods=["GET", "POST"])
def excel_to_pdf():
    if request.method == "POST":

        file = request.files.get("file")
        if not file:
            return "No file uploaded"

        input_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
        file.save(input_path)

        output_path = os.path.join(
            app.config["CONVERTED_FOLDER"],
            file.filename.replace(".xlsx", ".pdf")
        )

        convert_excel_to_pdf(input_path, output_path)

        output_filename = file.filename.replace(".xlsx", ".pdf")

        return send_from_directory(
            app.config["CONVERTED_FOLDER"],
            output_filename,
            as_attachment=True
        )

    return render_template("excel_to_pdf.html")

@app.route("/ppt-to-pdf", methods=["GET", "POST"])
def ppt_to_pdf():
    if request.method == "POST":

        file = request.files.get("file")
        if not file:
            return "No file uploaded"

        input_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
        file.save(input_path)

        output_path = os.path.join(
            app.config["CONVERTED_FOLDER"],
            file.filename.replace(".pptx", ".pdf")
        )

        convert_ppt_to_pdf(input_path, output_path)

        output_filename = file.filename.replace(".pptx", ".pdf")

        return send_from_directory(
            app.config["CONVERTED_FOLDER"],
            output_filename,
            as_attachment=True
        )

    return render_template("ppt_to_pdf.html")

from playwright.sync_api import sync_playwright

@app.route("/html-to-pdf", methods=["GET", "POST"])
def html_to_pdf():
    if request.method == "POST":

        url = request.form.get("url")
        html_code = request.form.get("html_code")

        output_path = os.path.join(
            app.config["CONVERTED_FOLDER"],
            "webpage.pdf"
        )

        try:
            with sync_playwright() as p:
                browser = p.chromium.launch(headless=True)
                page = browser.new_page()

                if url:
                    page.goto(url, wait_until="networkidle")
                elif html_code:
                    page.set_content(html_code)
                else:
                    return "No input provided"

                page.pdf(
                    path=output_path,
                    format="A4",
                    print_background=True
                )

                browser.close()

        except Exception as e:
            return f"Conversion failed: {str(e)}"

        return send_from_directory(
            app.config["CONVERTED_FOLDER"],
            "webpage.pdf",
            as_attachment=True
        )

    return render_template("html_to_pdf.html")

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from PyPDF2 import PdfReader, PdfWriter
import io

@app.route("/add-page-numbers", methods=["GET", "POST"])
def add_page_numbers():

    if request.method == "POST":

        file = request.files.get("file")
        position = request.form.get("position")  # top/bottom
        alignment = request.form.get("alignment")  # left/center/right

        if not file:
            return "No file uploaded"

        input_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
        file.save(input_path)

        reader = PdfReader(input_path)
        writer = PdfWriter()

        for i, page in enumerate(reader.pages):

            packet = io.BytesIO()
            can = canvas.Canvas(packet, pagesize=A4)

            page_number_text = f"Page {i+1}"

            width, height = A4

            y = height - 40 if position == "top" else 30

            if alignment == "left":
                x = 40
            elif alignment == "center":
                x = width / 2
            else:
                x = width - 100

            can.drawString(x, y, page_number_text)
            can.save()

            packet.seek(0)
            overlay = PdfReader(packet)
            page.merge_page(overlay.pages[0])
            writer.add_page(page)

        output_path = os.path.join(app.config["CONVERTED_FOLDER"], "numbered.pdf")

        with open(output_path, "wb") as f:
            writer.write(f)

        return send_from_directory(
            app.config["CONVERTED_FOLDER"],
            "numbered.pdf",
            as_attachment=True
        )

    return render_template("add_page_numbers.html")

from reportlab.pdfgen import canvas
from PyPDF2 import PdfReader, PdfWriter
from reportlab.lib.utils import ImageReader
import base64
import io

def generate_summary(text, max_length=150):
    """Generate a summary of the provided text using a simple extractive approach."""
    sentences = text.split('. ')
    if len(sentences) <= 3:
        return text
    summary_sentences = sentences[:3]
    return '. '.join(summary_sentences) + '.'

def safe_float(value, default=0):
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def safe_int(value, default=0):
    try:
        return int(value)
    except (TypeError, ValueError):
        return default

@app.route("/sign-pdf", methods=["GET", "POST"])
def sign_pdf():

    if request.method == "POST":

        file = request.files.get("file")
        signature_drawn = request.form.get("signature")
        signature_image_file = request.files.get("signature_image")

        text_value = request.form.get("text_value")
        text_x = safe_float(request.form.get("text_x"), 0)
        text_y = safe_float(request.form.get("text_y"), 0)
        text_size = safe_int(request.form.get("text_size"), 16)
        text_color = request.form.get("text_color", "#000000")

        page_number = safe_int(request.form.get("page_number"), 1) - 1
        pos_x = safe_float(request.form.get("pos_x"), 100)
        pos_y = safe_float(request.form.get("pos_y"), 100)

        input_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
        file.save(input_path)

        reader = PdfReader(input_path)
        writer = PdfWriter()

        signature_stream = None

        if signature_image_file and signature_image_file.filename != "":
            signature_stream = io.BytesIO(signature_image_file.read())
        elif signature_drawn:
            signature_drawn = signature_drawn.split(",")[1]
            signature_image = base64.b64decode(signature_drawn)
            signature_stream = io.BytesIO(signature_image)

        for i, page in enumerate(reader.pages):

            packet = io.BytesIO()
            can = canvas.Canvas(packet)

            draw_something = False

            if i == page_number and (signature_stream or text_value):

                if signature_stream:
                    can.drawImage(
                        ImageReader(signature_stream),
                        pos_x,
                        pos_y,
                        width=150,
                        height=50,
                        mask='auto'
                    )
                    draw_something = True

                if text_value:
                    r = int(text_color[1:3], 16)
                    g = int(text_color[3:5], 16)
                    b = int(text_color[5:7], 16)

                    can.setFillColorRGB(r/255, g/255, b/255)
                    can.setFont("Helvetica", text_size)
                    can.drawString(text_x, text_y, text_value)
                    draw_something = True

            can.save()
            packet.seek(0)

            if draw_something:
                overlay = PdfReader(packet)
                if len(overlay.pages) > 0:
                    page.merge_page(overlay.pages[0])

            writer.add_page(page)

        output_path = os.path.join(app.config["CONVERTED_FOLDER"], "signed.pdf")

        with open(output_path, "wb") as f:
            writer.write(f)

        return send_from_directory(
            app.config["CONVERTED_FOLDER"],
            "signed.pdf",
            as_attachment=True
        )

    return render_template("sign_pdf.html")

from utils.ai_tools import generate_summary, rewrite_text, translate_text

@app.route("/ai-summarise", methods=["GET", "POST"])
def ai_summarise():

    summary = None
    word_count = 0

    if request.method == "POST":

        file = request.files.get("file")
        text_input = request.form.get("text_input")
        mode = request.form.get("mode") or "standard"
        tone = request.form.get("tone") or "formal"

        extracted_text = ""

        # 🔹 If PDF uploaded
        if file and file.filename != "":
            from PyPDF2 import PdfReader
            reader = PdfReader(file)

            for page in reader.pages:
                extracted_text += page.extract_text() or ""

        # 🔹 If direct text input
        elif text_input:
            extracted_text = text_input

        # 🔹 Generate summary if text exists
        if extracted_text.strip():

            summary = generate_summary(
                extracted_text,
                mode,
                tone
            )

            # Word count (for display)
            word_count = len(summary.split())

    return render_template(
        "ai_summarise.html",
        summary=summary,
        word_count=word_count
    )

@app.route("/download-summary-pdf", methods=["POST"])
def download_summary_pdf():

    html_content = request.form.get("summary_html")

    output_path = os.path.join(
        app.config["CONVERTED_FOLDER"],
        "ai_summary.pdf"
    )

    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.set_content(html_content)
        page.pdf(path=output_path, format="A4")
        browser.close()

    return send_from_directory(
        app.config["CONVERTED_FOLDER"],
        "ai_summary.pdf",
        as_attachment=True
    )

@app.route("/insert-into-editor", methods=["POST"])
def insert_into_editor():

    summary_html = request.form.get("summary_html")

    filename = "AI_Summary.docx"
    output_path = os.path.join(app.config["CONVERTED_FOLDER"], filename)

    html_to_docx(summary_html, output_path)

    return redirect(url_for("editor", filename=filename))

@app.route("/ai-rewrite", methods=["POST"])
def ai_rewrite():

    summary_html = request.form.get("summary_html")
    action = request.form.get("action")

    if not summary_html:
        return {"error": "No content"}, 400

    rewritten = rewrite_text(summary_html, action)

    return {"result": rewritten}

@app.route("/ai-rewrite-tool", methods=["GET", "POST"])
def ai_rewrite_tool():

    rewritten = None

    if request.method == "POST":
        text_input = request.form.get("text_input")
        action = request.form.get("action") or "rewrite"

        if text_input:
            rewritten = rewrite_text(text_input, action)

    return render_template(
        "ai_rewrite.html",
        rewritten=rewritten
    )

from utils.ai_tools import translate_text, detect_language

@app.route("/ai-translate", methods=["GET", "POST"])
def ai_translate():

    translated = None
    detected_language = None
    original_text = None

    if request.method == "POST":

        file = request.files.get("file")
        text_input = request.form.get("text_input")
        target_language = request.form.get("language") or "Hindi"

        extracted_text = ""

        if file and file.filename != "":
            from PyPDF2 import PdfReader
            reader = PdfReader(file)
            for page in reader.pages:
                extracted_text += page.extract_text() or ""

        elif text_input:
            extracted_text = text_input

        if extracted_text.strip():
            original_text = extracted_text
            translated = translate_text(extracted_text, target_language)

            # 🔥 Detect language of translated result
            detected_language = detect_language(translated)

    return render_template(
        "ai_translate.html",
        translated=translated,
        original_text=original_text,
        detected_language=detected_language
    )

@app.route("/download-translation-pdf", methods=["POST"])
def download_translation_pdf():

    html_content = request.form.get("translated_html")

    output_path = os.path.join(
        app.config["CONVERTED_FOLDER"],
        "translation.pdf"
    )

    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.set_content(html_content)
        page.pdf(path=output_path, format="A4", print_background=True)
        browser.close()

    return send_from_directory(
        app.config["CONVERTED_FOLDER"],
        "translation.pdf",
        as_attachment=True
    )

from gtts import gTTS

@app.route("/download-translation-audio", methods=["POST"])
def download_translation_audio():

    text = request.form.get("translated_text")

    if not text:
        return "No text found"

    audio_path = os.path.join(
        app.config["CONVERTED_FOLDER"],
        "translation.mp3"
    )

    tts = gTTS(text=text, lang="en")
    tts.save(audio_path)

    return send_from_directory(
        app.config["CONVERTED_FOLDER"],
        "translation.mp3",
        as_attachment=True
    )

from flask import session, jsonify
from utils.ai_tools import chat_with_pdf

app.secret_key = "super_secret_key"  # required for session

@app.route("/ai-chat-pdf", methods=["GET", "POST"])
def ai_chat_pdf():

    if request.method == "POST":

        file = request.files.get("file")

        if file and file.filename != "":
            from PyPDF2 import PdfReader
            reader = PdfReader(file)

            pdf_text = ""
            for page in reader.pages:
                pdf_text += page.extract_text() or ""

            conn = get_connection()
            cursor = conn.cursor()

            cursor.execute("""
                INSERT INTO pdf_sessions (filename, extracted_text, created_at)
                VALUES (?, ?, ?)
            """, (file.filename, pdf_text, datetime.now().isoformat()))

            session_id = cursor.lastrowid

            conn.commit()
            conn.close()

            session["current_session_id"] = session_id

    # ✅ LOAD CHAT HISTORY FROM DATABASE HERE
    session_id = session.get("current_session_id")
    chat_history = []

    if session_id:
        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT role, message
            FROM chat_messages
            WHERE session_id = ?
            ORDER BY created_at ASC
        """, (session_id,))

        rows = cursor.fetchall()
        conn.close()

        chat_history = [
            {"role": row["role"], "content": row["message"]}
            for row in rows
        ]

    return render_template("ai_chat_pdf.html", chat_history=chat_history)

@app.route("/ask-pdf-question", methods=["POST"])
def ask_pdf_question():

    question = request.json.get("question")
    session_id = session.get("current_session_id")

    if not session_id:
        return jsonify({"error": "Upload PDF first"})

    conn = get_connection()
    cursor = conn.cursor()

    # Get PDF text
    cursor.execute(
        "SELECT extracted_text FROM pdf_sessions WHERE id = ?",
        (session_id,)
    )
    row = cursor.fetchone()

    if not row:
        conn.close()
        return jsonify({"error": "Session not found"})

    pdf_text = row["extracted_text"]

    # Generate AI answer
    result = chat_with_pdf(pdf_text, question)

    if isinstance(result, str):
        result = {
            "section": "Unknown",
            "snippet": "",
            "answer": result
        }

    # Save user message
    cursor.execute("""
        INSERT INTO chat_messages (session_id, role, message, created_at)
        VALUES (?, ?, ?, ?)
    """, (session_id, "user", question, datetime.now().isoformat()))

    # Save AI message
    cursor.execute("""
        INSERT INTO chat_messages (session_id, role, message, created_at)
        VALUES (?, ?, ?, ?)
    """, (session_id, "ai", result["answer"], datetime.now().isoformat()))

    # Save user question
    cursor.execute("""
    INSERT INTO chat_messages (session_id, role, message, created_at)
    VALUES (?, ?, ?, ?)
    """, (session_id, "user", question, datetime.now().isoformat()))

    # Save AI answer
    cursor.execute("""
    INSERT INTO chat_messages (session_id, role, message, created_at)
    VALUES (?, ?, ?, ?)
    """, (session_id, "ai", result["answer"], datetime.now().isoformat()))

    conn.commit()
    conn.close()

    return jsonify(result)

@app.route("/ask-global-ai", methods=["POST"])
def ask_global_ai():

    question = request.json.get("question")

    if not question:
        return jsonify({"answer": "Please enter a question."})

    # Simple AI assistant mode
    result = chat_with_pdf(
        "You are a helpful AI assistant for ME & PDFs platform.",
        question
    )

    if isinstance(result, dict):
        return jsonify(result)
    else:
        return jsonify({"answer": result})
    
from PyPDF2 import PdfReader, PdfWriter

@app.route("/unlock-pdf", methods=["GET", "POST"])
def unlock_pdf():

    if request.method == "POST":

        file = request.files.get("file")
        password = request.form.get("password")

        if not file:
            return "No file uploaded"

        input_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
        output_filename = "unlocked.pdf"
        output_path = os.path.join(app.config["CONVERTED_FOLDER"], output_filename)

        file.save(input_path)

        reader = PdfReader(input_path)

        if reader.is_encrypted:
            try:
                reader.decrypt(password)
            except:
                return "Incorrect password"

        writer = PdfWriter()

        for page in reader.pages:
            writer.add_page(page)

        with open(output_path, "wb") as f:
            writer.write(f)

        return send_from_directory(
            app.config["CONVERTED_FOLDER"],
            output_filename,
            as_attachment=True
        )

    return render_template("unlock_pdf.html", error="Incorrect password")

import pikepdf

@app.route("/protect-pdf", methods=["GET", "POST"])
def protect_pdf():

    if request.method == "POST":

        file = request.files.get("file")
        password = request.form.get("password")

        if not file or not password:
            return "File and password required"

        input_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
        output_filename = "protected.pdf"
        output_path = os.path.join(app.config["CONVERTED_FOLDER"], output_filename)

        file.save(input_path)

        try:
            with pikepdf.open(input_path) as pdf:

                pdf.save(
                    output_path,
                    encryption=pikepdf.Encryption(
                        user=password,
                        owner=password,
                        R=4  # AES 128-bit encryption
                    )
                )

        except Exception as e:
            return f"Error: {str(e)}"

        return send_from_directory(
            app.config["CONVERTED_FOLDER"],
            output_filename,
            as_attachment=True
        )

    return render_template("protect_pdf.html")

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import Color
from PyPDF2 import PdfReader, PdfWriter
import io

@app.route("/watermark-pdf", methods=["GET", "POST"])
def watermark_pdf():

    if request.method == "POST":

        file = request.files.get("file")
        watermark_text = request.form.get("watermark_text")
        position = request.form.get("position")
        opacity = float(request.form.get("opacity", 0.3))

        if not file or not watermark_text:
            return "File and watermark text required"

        input_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
        output_filename = "watermarked.pdf"
        output_path = os.path.join(app.config["CONVERTED_FOLDER"], output_filename)

        file.save(input_path)

        reader = PdfReader(input_path)
        writer = PdfWriter()

        for page in reader.pages:

            packet = io.BytesIO()
            page_width = float(page.mediabox.width)
            page_height = float(page.mediabox.height)

            can = canvas.Canvas(packet, pagesize=(page_width, page_height))

            can.setFont("Helvetica", 40)
            can.setFillColor(Color(0.8, 0.8, 0.8, alpha=opacity))

            # Position logic
            if position == "center":
                can.drawCentredString(page_width / 2, page_height / 2, watermark_text)

            elif position == "top-left":
                can.drawString(50, page_height - 50, watermark_text)

            elif position == "bottom-right":
                can.drawRightString(page_width - 50, 50, watermark_text)

            else:
                can.drawCentredString(page_width / 2, page_height / 2, watermark_text)

            can.save()

            packet.seek(0)
            watermark_pdf = PdfReader(packet)
            watermark_page = watermark_pdf.pages[0]

            page.merge_page(watermark_page)
            writer.add_page(page)

        with open(output_path, "wb") as f:
            writer.write(f)

        return send_from_directory(
            app.config["CONVERTED_FOLDER"],
            output_filename,
            as_attachment=True
        )

    return render_template("watermark_pdf.html")


if __name__ == "__main__":
    app.run(debug=True)